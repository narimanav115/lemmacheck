#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LemmaCheck - Full-text search application with lemmatization support
Supports Russian (pymorphy3), English (nltk) and Kazakh (built-in stemmer) lemmatization
Uses PyQt6 for GUI
"""

import sys
import os
import re
import json
import csv
import math
import subprocess
import platform
from datetime import datetime
from collections import defaultdict
from typing import Dict, List, Tuple, Set, Optional
from pathlib import Path

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLineEdit, QListWidget, QTextEdit, QLabel,
    QFileDialog, QMessageBox, QProgressBar, QGroupBox, QSplitter,
    QAbstractItemView, QListWidgetItem, QFrame, QCheckBox,
    QComboBox, QTableWidget, QTableWidgetItem, QHeaderView,
    QTabWidget, QDialog, QRadioButton, QButtonGroup
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtGui import QFont, QTextCharFormat, QColor, QTextCursor, QBrush, QPalette

import chardet

# Document parsers
try:
    from docx import Document
except ImportError:
    Document = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

# Excel support
try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import xlrd
except ImportError:
    xlrd = None

# Windows COM support for .doc files
win32com_client = None
if platform.system() == 'Windows':
    try:
        import win32com.client as win32com_client
    except ImportError:
        pass

# Lemmatization
try:
    import pymorphy3
    morph = pymorphy3.MorphAnalyzer()
except ImportError:
    morph = None

try:
    import nltk
    from nltk.stem import WordNetLemmatizer
    from nltk.corpus import wordnet
    try:
        nltk.data.find('corpora/wordnet')
    except LookupError:
        nltk.download('wordnet', quiet=True)
    try:
        nltk.data.find('taggers/averaged_perceptron_tagger')
    except LookupError:
        nltk.download('averaged_perceptron_tagger', quiet=True)
    try:
        nltk.data.find('taggers/averaged_perceptron_tagger_eng')
    except LookupError:
        nltk.download('averaged_perceptron_tagger_eng', quiet=True)
    lemmatizer = WordNetLemmatizer()
except ImportError:
    lemmatizer = None


# ==================== Kazakh Stemmer ====================

# Kazakh-specific Cyrillic characters (not present in Russian)
KAZAKH_SPECIFIC_CHARS = set('ӘәҒғҚқҢңӨөҰұҮүҺһІі')

# Regex character class for all Cyrillic including Kazakh-specific letters
CYRILLIC_CHARS = r'а-яёА-ЯЁӘәҒғҚқҢңӨөҰұҮүҺһІі'
# Full word pattern: Cyrillic + Latin + hyphenated compounds
WORD_PATTERN = r'[' + CYRILLIC_CHARS + r'a-zA-Z]+(?:-[' + CYRILLIC_CHARS + r'a-zA-Z]+)*'


class KazakhStemmer:
    """Rule-based suffix stemmer for Kazakh language.
    
    Kazakh is agglutinative — words are formed by appending suffixes to stems.
    This stemmer strips common inflectional and derivational suffixes.
    """

    # Plural suffixes
    _PLURAL = ['лар', 'лер', 'дар', 'дер', 'тар', 'тер']

    # Case suffixes (genitive, dative, accusative, locative, ablative, instrumental)
    _CASE = [
        'ның', 'нің',           # genitive
        'ға', 'ге', 'қа', 'ке', # dative
        'ны', 'ні', 'ды', 'ді', 'ты', 'ті',  # accusative
        'да', 'де', 'та', 'те', # locative
        'дан', 'ден', 'тан', 'тен', 'нан', 'нен',  # ablative
        'мен', 'бен', 'пен',   # instrumental
    ]

    # Possessive suffixes
    _POSSESSIVE = [
        'ым', 'ім', 'м',       # 1st person sing
        'ың', 'ің', 'ң',       # 2nd person sing
        'ы', 'і', 'сы', 'сі',  # 3rd person sing
        'мыз', 'міз',          # 1st person plur
        'ңыз', 'ңіз',          # 2nd person formal
        'дары', 'дері', 'тары', 'тері', 'лары', 'лері',  # 3rd person plur
    ]

    # Verbal suffixes (common tense/mood markers)
    _VERBAL = [
        'ған', 'ген', 'қан', 'кен',   # past participle
        'атын', 'етін',                # habitual
        'йтын', 'йтін',               # habitual
        'ушы', 'уші',                  # agent noun
        'лық', 'лік', 'дық', 'дік', 'тық', 'тік',  # abstract noun
        'шы', 'ші',                    # agent/profession
        'сыз', 'сіз',                  # privative (without)
    ]

    def __init__(self):
        # Build suffix list sorted by length (longest first) for greedy matching
        all_suffixes = (self._PLURAL + self._CASE + self._POSSESSIVE + self._VERBAL)
        self._suffixes = sorted(set(all_suffixes), key=len, reverse=True)

    def stem(self, word: str) -> str:
        """Strip Kazakh suffixes from a word to approximate its lemma."""
        word = word.lower().strip()
        if len(word) < 3:
            return word

        # Apply up to 3 rounds of suffix stripping (Kazakh can stack multiple suffixes)
        for _ in range(3):
            stripped = False
            for suffix in self._suffixes:
                if word.endswith(suffix) and len(word) - len(suffix) >= 2:
                    word = word[:-len(suffix)]
                    stripped = True
                    break
            if not stripped:
                break

        return word


# Global Kazakh stemmer instance
kaz_stemmer = KazakhStemmer()


# ==================== Document Parsers ====================

def extract_text_from_docx(filepath: str) -> str:
    if Document is None:
        raise ImportError("python-docx не установлен")
    doc = Document(filepath)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return '\n'.join(paragraphs)


def extract_text_from_pdf(filepath: str) -> str:
    if fitz is None:
        raise ImportError("PyMuPDF не установлен")
    text_parts = []
    with fitz.open(filepath) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return '\n'.join(text_parts)


def extract_text_from_txt(filepath: str) -> str:
    with open(filepath, 'rb') as f:
        raw_data = f.read()
    detected = chardet.detect(raw_data)
    encoding = detected.get('encoding', 'utf-8')
    for enc in [encoding, 'utf-8', 'cp1251', 'cp1252', 'latin-1']:
        if enc:
            try:
                return raw_data.decode(enc)
            except (UnicodeDecodeError, LookupError):
                continue
    return raw_data.decode('utf-8', errors='ignore')


def extract_text(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.docx':
        return extract_text_from_docx(filepath)
    elif ext == '.doc':
        return extract_text_from_doc(filepath)
    elif ext == '.pdf':
        return extract_text_from_pdf(filepath)
    elif ext == '.txt':
        return extract_text_from_txt(filepath)
    elif ext == '.xlsx':
        return extract_text_from_xlsx(filepath)
    elif ext == '.xls':
        return extract_text_from_xls(filepath)
    raise ValueError(f"Неподдерживаемый формат: {ext}")


def extract_text_from_doc(filepath: str) -> str:
    """Extract text from old .doc format using platform-specific methods"""
    import subprocess
    import tempfile
    
    system = platform.system()
    
    # Windows: use pywin32 COM automation
    if system == 'Windows':
        if win32com_client is not None:
            try:
                word = win32com_client.Dispatch('Word.Application')
                word.Visible = False
                doc = word.Documents.Open(os.path.abspath(filepath))
                text = doc.Content.Text
                doc.Close(False)
                word.Quit()
                return text
            except Exception:
                pass
    
    # macOS: use textutil (built-in)
    if system == 'Darwin':
        try:
            with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
                tmp_path = tmp.name
            subprocess.run(['textutil', '-convert', 'txt', '-output', tmp_path, filepath], 
                          check=True, capture_output=True)
            with open(tmp_path, 'r', encoding='utf-8') as f:
                text = f.read()
            os.unlink(tmp_path)
            return text
        except Exception:
            pass
    
    # Cross-platform fallback: antiword
    try:
        result = subprocess.run(['antiword', filepath], capture_output=True, text=True, check=True)
        return result.stdout
    except Exception:
        pass
    
    # Platform-specific error messages
    if system == 'Windows':
        raise ImportError("Не удалось прочитать .doc файл. Установите pywin32: pip install pywin32")
    elif system == 'Darwin':
        raise ImportError("Не удалось прочитать .doc файл с помощью textutil.")
    else:
        raise ImportError("Не удалось прочитать .doc файл. Установите antiword.")


def extract_text_from_xlsx(filepath: str) -> str:
    """Extract text from .xlsx files"""
    if openpyxl is None:
        raise ImportError("openpyxl не установлен. Установите: pip install openpyxl")
    
    wb = openpyxl.load_workbook(filepath, data_only=True)
    texts = []
    
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            row_texts = []
            for cell in row:
                if cell.value is not None:
                    row_texts.append(str(cell.value))
            if row_texts:
                texts.append(' '.join(row_texts))
    
    return '\n'.join(texts)


def extract_text_from_xls(filepath: str) -> str:
    """Extract text from old .xls files"""
    if xlrd is None:
        raise ImportError("xlrd не установлен. Установите: pip install xlrd")
    
    wb = xlrd.open_workbook(filepath)
    texts = []
    
    for sheet in wb.sheets():
        for row_idx in range(sheet.nrows):
            row_texts = []
            for col_idx in range(sheet.ncols):
                cell_value = sheet.cell_value(row_idx, col_idx)
                if cell_value:
                    row_texts.append(str(cell_value))
            if row_texts:
                texts.append(' '.join(row_texts))
    
    return '\n'.join(texts)


# ==================== Search Engine ====================

class LemmaSearchEngine:
    def __init__(self):
        self.documents: Dict[str, dict] = {}
        self.inverted_index: Dict[str, Dict[str, List[int]]] = defaultdict(lambda: defaultdict(list))
        self.doc_frequency: Dict[str, int] = defaultdict(int)
        self.total_docs: int = 0
        self.total_words: int = 0
        self._lemma_cache: Dict[str, str] = {}
        self.force_kazakh: bool = False

    def _is_cyrillic(self, word: str) -> bool:
        return bool(re.search(r'[а-яёА-ЯЁӘәҒғҚқҢңӨөҰұҮүҺһІі]', word))

    def _is_kazakh(self, word: str) -> bool:
        """Detect Kazakh by presence of Kazakh-specific Cyrillic characters."""
        return bool(KAZAKH_SPECIFIC_CHARS & set(word))

    def _is_latin(self, word: str) -> bool:
        return bool(re.search(r'[a-zA-Z]', word))

    def _get_wordnet_pos(self, word: str) -> str:
        try:
            tag = nltk.pos_tag([word])[0][1][0].upper()
            tag_dict = {'J': wordnet.ADJ, 'N': wordnet.NOUN, 'V': wordnet.VERB, 'R': wordnet.ADV}
            return tag_dict.get(tag, wordnet.NOUN)
        except:
            return wordnet.NOUN

    def _lemmatize_word(self, word: str) -> str:
        word_lower = word.lower().strip()
        if not word_lower or len(word_lower) < 2:
            return word_lower
        if word_lower in self._lemma_cache:
            return self._lemma_cache[word_lower]
        lemma = word_lower
        if self._is_cyrillic(word_lower):
            if self.force_kazakh or self._is_kazakh(word_lower):
                # Kazakh mode (forced or detected) — use Kazakh stemmer
                lemma = kaz_stemmer.stem(word_lower)
            elif morph:
                # Standard Russian Cyrillic — use pymorphy3
                parsed = morph.parse(word_lower)
                if parsed:
                    lemma = parsed[0].normal_form
        elif self._is_latin(word_lower) and lemmatizer:
            pos = self._get_wordnet_pos(word_lower)
            lemma = lemmatizer.lemmatize(word_lower, pos)
        self._lemma_cache[word_lower] = lemma
        return lemma

    def lemmatize(self, text: str) -> List[Tuple[str, str, int]]:
        """Tokenize and lemmatize text, supporting hyphenated words"""
        words = []
        # Match words including hyphenated compounds (e.g., "Three-cycle", "научно-технический", "ғылыми-техникалық")
        for match in re.finditer(WORD_PATTERN, text):
            word = match.group()
            position = match.start()
            
            # Handle hyphenated words
            if '-' in word:
                # Index the full compound word
                full_lemma = self._lemmatize_compound(word)
                words.append((word, full_lemma, position))
                
                # Also index individual parts for broader matching
                parts = word.split('-')
                offset = 0
                for part in parts:
                    if part:
                        part_lemma = self._lemmatize_word(part)
                        words.append((part, part_lemma, position + offset))
                    offset += len(part) + 1  # +1 for hyphen
            else:
                lemma = self._lemmatize_word(word)
                words.append((word, lemma, position))
        return words

    def _lemmatize_compound(self, word: str) -> str:
        """Lemmatize a hyphenated compound word"""
        parts = word.split('-')
        lemmatized_parts = [self._lemmatize_word(part) for part in parts if part]
        return '-'.join(lemmatized_parts)

    def add_document(self, doc_id: str, text: str, filename: str) -> int:
        if doc_id in self.documents:
            self.remove_document(doc_id)
        words = self.lemmatize(text)
        doc_lemmas: Set[str] = set()
        for original, lemma, position in words:
            if lemma:
                self.inverted_index[lemma][doc_id].append(position)
                doc_lemmas.add(lemma)
        for lemma in doc_lemmas:
            self.doc_frequency[lemma] += 1
        self.documents[doc_id] = {'filename': filename, 'text': text, 'word_count': len(words)}
        self.total_docs += 1
        self.total_words += len(words)
        return len(words)

    def remove_document(self, doc_id: str):
        if doc_id not in self.documents:
            return
        lemmas_to_check = [l for l, docs in self.inverted_index.items() if doc_id in docs]
        for lemma in lemmas_to_check:
            del self.inverted_index[lemma][doc_id]
            self.doc_frequency[lemma] -= 1
            if self.doc_frequency[lemma] <= 0:
                del self.doc_frequency[lemma]
            if not self.inverted_index[lemma]:
                del self.inverted_index[lemma]
        self.total_words -= self.documents[doc_id]['word_count']
        self.total_docs -= 1
        del self.documents[doc_id]

    def _calculate_tfidf(self, lemma: str, doc_id: str) -> float:
        if lemma not in self.inverted_index or doc_id not in self.inverted_index[lemma]:
            return 0.0
        tf = len(self.inverted_index[lemma][doc_id])
        doc_word_count = self.documents[doc_id]['word_count']
        if doc_word_count > 0:
            tf = tf / doc_word_count
        df = self.doc_frequency.get(lemma, 0)
        idf = math.log(self.total_docs / df) + 1 if df > 0 and self.total_docs > 0 else 1
        return tf * idf

    def search(self, query: str) -> List[Tuple[str, int, str, List[int]]]:
        """Search for exact phrase (words in sequence) in documents"""
        if not query.strip():
            return []
        
        query_words = self.lemmatize(query)
        query_lemmas = [lemma for _, lemma, _ in query_words if lemma]
        if not query_lemmas:
            return []
        
        query_lemmas_set = set(query_lemmas)
        
        # Find candidate documents that have ALL query lemmas
        candidate_docs: Set[str] = set()
        for lemma in query_lemmas_set:
            if lemma in self.inverted_index:
                if not candidate_docs:
                    candidate_docs = set(self.inverted_index[lemma].keys())
                else:
                    candidate_docs &= set(self.inverted_index[lemma].keys())
        
        if not candidate_docs:
            return []
        
        results = []
        
        # For single word queries, just count occurrences
        if len(query_lemmas) == 1:
            lemma = query_lemmas[0]
            for doc_id in candidate_docs:
                positions = self.inverted_index[lemma][doc_id]
                count = len(positions)
                results.append((doc_id, count, self.documents[doc_id]['filename'], positions))
        else:
            # For multi-word queries, find exact phrase matches
            for doc_id in candidate_docs:
                phrase_positions = self._find_phrase_in_document(doc_id, query_lemmas)
                if phrase_positions:
                    count = len(phrase_positions)
                    results.append((doc_id, count, self.documents[doc_id]['filename'], phrase_positions))
        
        results.sort(key=lambda x: x[1], reverse=True)
        return results

    def _find_phrase_in_document(self, doc_id: str, query_lemmas: List[str], max_distance: int = 15) -> List[int]:
        """Find all occurrences of a phrase in a document.
        Returns list of starting positions where the phrase was found.
        max_distance is the maximum character distance between consecutive words."""
        if doc_id not in self.documents:
            return []
        
        # Get positions for each lemma in the query
        lemma_positions = []
        for lemma in query_lemmas:
            if lemma in self.inverted_index and doc_id in self.inverted_index[lemma]:
                lemma_positions.append(sorted(self.inverted_index[lemma][doc_id]))
            else:
                return []  # If any lemma is missing, no phrase match possible
        
        if not lemma_positions or not lemma_positions[0]:
            return []
        
        phrase_starts = []
        first_positions = lemma_positions[0]
        
        for start_pos in first_positions:
            current_pos = start_pos
            matched = True
            
            for i in range(1, len(lemma_positions)):
                next_positions = lemma_positions[i]
                found_next = False
                
                # Find the next lemma position that comes after current_pos within max_distance
                for pos in next_positions:
                    if current_pos < pos <= current_pos + max_distance:
                        current_pos = pos
                        found_next = True
                        break
                
                if not found_next:
                    matched = False
                    break
            
            if matched:
                phrase_starts.append(start_pos)
        
        return phrase_starts

    def get_query_lemmas(self, query: str) -> List[str]:
        """Get ordered list of lemmas from query"""
        return [lemma for _, lemma, _ in self.lemmatize(query) if lemma]

    def _parse_phrases(self, query: str) -> List[List[str]]:
        """Parse query into phrases (multi-word sequences)"""
        phrases = []
        
        # Handle quoted phrases first
        quoted_pattern = r'"([^"]+)"|\'([^\']+)\''
        for match in re.finditer(quoted_pattern, query):
            phrase_text = match.group(1) or match.group(2)
            phrase_lemmas = [lemma for _, lemma, _ in self.lemmatize(phrase_text) if lemma]
            if len(phrase_lemmas) >= 2:
                phrases.append(phrase_lemmas)
        
        # Remove quoted parts from query
        remaining_query = re.sub(quoted_pattern, ' ', query)
        
        # Split by common phrase delimiters: semicolons, commas
        parts = re.split(r'[;,]', remaining_query)
        
        for part in parts:
            part = part.strip()
            if not part:
                continue
            
            # Each part with 2+ words is treated as a potential phrase
            part_lemmas = [lemma for _, lemma, _ in self.lemmatize(part) if lemma]
            if len(part_lemmas) >= 2:
                phrases.append(part_lemmas)
        
        return phrases

    def _calculate_phrase_boost(self, doc_id: str, phrases: List[List[str]], max_distance: int = 10) -> float:
        """Calculate boost factor based on how well phrases match in the document"""
        if doc_id not in self.documents:
            return 0.0
        
        total_boost = 0.0
        
        for phrase_lemmas in phrases:
            if len(phrase_lemmas) < 2:
                continue
            
            # Get positions for each lemma in the phrase
            lemma_positions = []
            for lemma in phrase_lemmas:
                if lemma in self.inverted_index and doc_id in self.inverted_index[lemma]:
                    lemma_positions.append(self.inverted_index[lemma][doc_id])
                else:
                    lemma_positions.append([])
            
            # Check if all lemmas are present
            if any(len(pos) == 0 for pos in lemma_positions):
                continue
            
            # Find sequences where words appear close together
            phrase_matches = self._find_phrase_matches(lemma_positions, max_distance)
            
            if phrase_matches > 0:
                # Boost based on number of phrase matches and phrase length
                phrase_boost = phrase_matches * (len(phrase_lemmas) ** 1.5) * 0.5
                total_boost += phrase_boost
        
        return min(total_boost, 5.0)  # Cap the boost

    def _find_phrase_matches(self, lemma_positions: List[List[int]], max_distance: int) -> int:
        """Find how many times words appear in sequence within max_distance"""
        if not lemma_positions or not lemma_positions[0]:
            return 0
        
        matches = 0
        first_positions = lemma_positions[0]
        
        for start_pos in first_positions:
            current_pos = start_pos
            matched = True
            
            for i in range(1, len(lemma_positions)):
                # Find next lemma position that's after current_pos but within max_distance
                next_positions = lemma_positions[i]
                found_next = False
                
                for pos in next_positions:
                    if current_pos < pos <= current_pos + max_distance:
                        current_pos = pos
                        found_next = True
                        break
                
                if not found_next:
                    matched = False
                    break
            
            if matched:
                matches += 1
        
        return matches

    def get_context(self, doc_id: str, positions: List[int], context_size: int = 50) -> str:
        if doc_id not in self.documents:
            return ""
        text = self.documents[doc_id]['text']
        if not positions:
            return text[:200] + "..." if len(text) > 200 else text
        contexts = []
        used_ranges = []
        for pos in positions[:3]:
            start = max(0, pos - context_size)
            end = min(len(text), pos + context_size)
            overlap = any(start < r_end and end > r_start for r_start, r_end in used_ranges)
            if not overlap:
                context = text[start:end]
                if start > 0:
                    context = "..." + context
                if end < len(text):
                    context = context + "..."
                contexts.append(context)
                used_ranges.append((start, end))
        return " | ".join(contexts)

    def get_sentences_with_matches(self, doc_id: str, positions: List[int]) -> List[Tuple[int, str]]:
        """Get sentences containing matches with 1 sentence before and after.
        Returns list of (match_number, context_text) tuples."""
        if doc_id not in self.documents:
            return []
        
        text = self.documents[doc_id]['text']
        if not positions:
            return []
        
        # Split text into sentences
        # Pattern matches sentence endings: . ! ? followed by space or newline
        sentence_pattern = r'[^.!?]*[.!?]+(?:\s|$)|[^.!?\n]+$'
        sentences = []
        for match in re.finditer(sentence_pattern, text):
            sent_text = match.group().strip()
            if sent_text:
                sentences.append((match.start(), match.end(), sent_text))
        
        if not sentences:
            return [(1, text[:500] + "..." if len(text) > 500 else text)]
        
        # Find which sentences contain matches
        matched_contexts = []
        used_positions = set()
        
        for match_idx, pos in enumerate(positions):
            if pos in used_positions:
                continue
                
            # Find sentence containing this position
            sent_idx = None
            for idx, (start, end, sent_text) in enumerate(sentences):
                if start <= pos < end:
                    sent_idx = idx
                    break
            
            if sent_idx is None:
                continue
            
            # Mark this position as used
            used_positions.add(pos)
            
            # Get 1 sentence before, the match sentence, and 1 sentence after
            context_parts = []
            
            # Previous sentence
            if sent_idx > 0:
                context_parts.append(sentences[sent_idx - 1][2])
            
            # Current sentence (with match)
            context_parts.append(sentences[sent_idx][2])
            
            # Next sentence
            if sent_idx < len(sentences) - 1:
                context_parts.append(sentences[sent_idx + 1][2])
            
            context = ' '.join(context_parts)
            
            # Check if this context overlaps with already added ones
            is_duplicate = False
            for _, existing_context in matched_contexts:
                if sentences[sent_idx][2] in existing_context:
                    is_duplicate = True
                    break
            
            if not is_duplicate:
                matched_contexts.append((match_idx + 1, context))
        
        return matched_contexts

    def get_kwic_concordance(self, doc_id: str, positions: List[int],
                              query_lemmas: List[str], context_words: int = 5,
                              context_type: str = 'words') -> List[Tuple[str, str, str, str, int]]:
        """Build KWIC concordance lines for a document.
        Returns list of (filename, left_context, keyword, right_context) tuples."""
        if doc_id not in self.documents:
            return []
        text = self.documents[doc_id]['text']
        filename = self.documents[doc_id]['filename']
        query_len = len(query_lemmas)
        results: List[Tuple[str, str, str, str, int]] = []

        # Build a word list with positions
        word_matches = list(re.finditer(WORD_PATTERN, text))
        if not word_matches:
            return []

        # Map char positions to word indices for fast lookup
        pos_to_widx: Dict[int, int] = {m.start(): i for i, m in enumerate(word_matches)}

        for pos in positions:
            # Find the word index closest to this char position
            widx = pos_to_widx.get(pos)
            if widx is None:
                # Find nearest
                best = None
                for start, idx in pos_to_widx.items():
                    if best is None or abs(start - pos) < abs(best - pos):
                        best = start
                if best is not None:
                    widx = pos_to_widx[best]
                else:
                    continue

            if context_type == 'sentence':
                # Sentence-based context
                kw_start = word_matches[widx].start()
                kw_end_idx = min(widx + query_len - 1, len(word_matches) - 1)
                kw_end = word_matches[kw_end_idx].end()
                keyword = text[kw_start:kw_end]

                # Find sentence boundaries
                sent_pattern = r'[^.!?]*[.!?]+(?:\s|$)|[^.!?\n]+$'
                sentences = [(m.start(), m.end()) for m in re.finditer(sent_pattern, text)]
                sent_start = 0
                sent_end = len(text)
                for s_start, s_end in sentences:
                    if s_start <= kw_start < s_end:
                        sent_start = s_start
                        sent_end = s_end
                        break

                left = text[sent_start:kw_start].strip()
                right = text[kw_end:sent_end].strip()
            else:
                # Word-based context (±N words)
                kw_end_idx = min(widx + query_len - 1, len(word_matches) - 1)
                keyword = text[word_matches[widx].start():word_matches[kw_end_idx].end()]

                left_start_idx = max(0, widx - context_words)
                right_end_idx = min(len(word_matches) - 1, kw_end_idx + context_words)

                if left_start_idx < widx:
                    left = text[word_matches[left_start_idx].start():word_matches[widx].start()].strip()
                else:
                    left = ''
                if kw_end_idx < right_end_idx:
                    right = text[word_matches[kw_end_idx].end():word_matches[right_end_idx].end()].strip()
                else:
                    right = ''

            results.append((filename, left, keyword, right, pos))

        return results

    def save_index(self, filepath: str):
        data = {
            'documents': self.documents,
            'inverted_index': {k: dict(v) for k, v in self.inverted_index.items()},
            'doc_frequency': dict(self.doc_frequency),
            'total_docs': self.total_docs,
            'total_words': self.total_words,
            'force_kazakh': self.force_kazakh
        }
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def load_index(self, filepath: str):
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        self.documents = data['documents']
        self.inverted_index = defaultdict(lambda: defaultdict(list))
        for lemma, docs in data['inverted_index'].items():
            for doc_id, positions in docs.items():
                self.inverted_index[lemma][doc_id] = positions
        self.doc_frequency = defaultdict(int, data['doc_frequency'])
        self.total_docs = data['total_docs']
        self.total_words = data['total_words']
        self.force_kazakh = data.get('force_kazakh', False)
        self._lemma_cache.clear()


# ==================== Indexing Thread ====================

class IndexingThread(QThread):
    progress = pyqtSignal(str, int, int)
    finished_file = pyqtSignal(str, str)
    error = pyqtSignal(str)
    done = pyqtSignal()

    def __init__(self, files: List[str], engine: LemmaSearchEngine):
        super().__init__()
        self.files = files
        self.engine = engine
        self.doc_paths: Dict[str, str] = {}
        self.cancelled = False

    def run(self):
        for i, filepath in enumerate(self.files):
            if self.cancelled:
                break
            filename = os.path.basename(filepath)
            self.progress.emit(f"Индексация: {filename}", i + 1, len(self.files))
            try:
                text = extract_text(filepath)
                self.engine.add_document(filepath, text, filename)
                self.doc_paths[filepath] = filepath
                self.finished_file.emit(filepath, filename)
            except Exception as e:
                self.error.emit(f"{filename}: {str(e)}")
        self.done.emit()


# ==================== Search Thread ====================

class SearchThread(QThread):
    """Thread for performing search and collecting results, optionally KWIC concordance."""
    progress = pyqtSignal(str)
    result_ready = pyqtSignal(object, str)  # (results, query)
    kwic_ready = pyqtSignal(object)  # kwic_rows list
    done = pyqtSignal()

    def __init__(self, engine: LemmaSearchEngine, query: str,
                 kwic_enabled: bool = False, kwic_context_type: str = 'words',
                 kwic_context_words: int = 5, kwic_filter: str = ''):
        super().__init__()
        self.engine = engine
        self.query = query
        self.kwic_enabled = kwic_enabled
        self.kwic_context_type = kwic_context_type
        self.kwic_context_words = kwic_context_words
        self.kwic_filter = kwic_filter
        self.results = []

    def run(self):
        self.progress.emit("Поиск совпадений...")
        self.results = self.engine.search(self.query)
        self.result_ready.emit(self.results, self.query)

        if self.kwic_enabled and self.results:
            self.progress.emit("Построение KWIC-конкорданса...")
            query_lemmas = self.engine.get_query_lemmas(self.query)
            all_rows: list = []
            for doc_id, count, filename, positions in self.results:
                rows = self.engine.get_kwic_concordance(
                    doc_id, positions, query_lemmas,
                    self.kwic_context_words, self.kwic_context_type
                )
                all_rows.extend(rows)

            # Apply context filter
            if self.kwic_filter:
                filter_lower = self.kwic_filter.lower()
                filter_lemma = self.engine._lemmatize_word(self.kwic_filter)
                filtered: list = []
                for row in all_rows:
                    fname, left, kw, right, pos = row
                    combined = f"{left} {right}".lower()
                    if filter_lower in combined:
                        filtered.append(row)
                        continue
                    context_words_list = re.findall(WORD_PATTERN, combined)
                    context_lemmas = [self.engine._lemmatize_word(w) for w in context_words_list]
                    if filter_lemma in context_lemmas:
                        filtered.append(row)
                all_rows = filtered

            self.kwic_ready.emit(all_rows)

        self.done.emit()


# ==================== Export Dialog ====================

class ExportDialog(QDialog):
    """Dialog for choosing export type and format."""

    def __init__(self, parent=None, has_kwic: bool = False):
        super().__init__(parent)
        self.setWindowTitle("Экспорт данных")
        self.setMinimumWidth(400)

        layout = QVBoxLayout(self)

        # Export type
        type_group = QGroupBox("Тип экспорта")
        type_layout = QVBoxLayout(type_group)

        self.type_button_group = QButtonGroup(self)
        self.radio_results = QRadioButton("Таблица результатов (частоты, IP10K, IPM, TF-IDF)")
        self.radio_concordance = QRadioButton("Конкорданс (KWIC-таблица)")
        self.radio_summary = QRadioButton("Сводный отчёт")

        self.radio_results.setChecked(True)
        self.radio_concordance.setEnabled(has_kwic)
        if not has_kwic:
            self.radio_concordance.setToolTip("Сначала включите KWIC и выполните поиск")

        self.type_button_group.addButton(self.radio_results, 0)
        self.type_button_group.addButton(self.radio_concordance, 1)
        self.type_button_group.addButton(self.radio_summary, 2)

        type_layout.addWidget(self.radio_results)
        type_layout.addWidget(self.radio_concordance)
        type_layout.addWidget(self.radio_summary)
        layout.addWidget(type_group)

        # Format
        format_group = QGroupBox("Формат файла")
        format_layout = QHBoxLayout(format_group)

        self.format_button_group = QButtonGroup(self)
        self.radio_csv = QRadioButton("CSV")
        self.radio_xlsx = QRadioButton("XLSX")
        self.radio_csv.setChecked(True)

        if openpyxl is None:
            self.radio_xlsx.setEnabled(False)
            self.radio_xlsx.setToolTip("openpyxl не установлен (pip install openpyxl)")

        self.format_button_group.addButton(self.radio_csv, 0)
        self.format_button_group.addButton(self.radio_xlsx, 1)

        format_layout.addWidget(self.radio_csv)
        format_layout.addWidget(self.radio_xlsx)
        layout.addWidget(format_group)

        # Buttons
        btn_layout = QHBoxLayout()
        btn_save = QPushButton("Сохранить")
        btn_save.setStyleSheet("background-color: #ec407a; color: white;")
        btn_save.clicked.connect(self.accept)
        btn_cancel = QPushButton("Отмена")
        btn_cancel.clicked.connect(self.reject)
        btn_layout.addStretch()
        btn_layout.addWidget(btn_save)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)

    def get_export_type(self) -> int:
        """0 = results table, 1 = concordance, 2 = summary"""
        return self.type_button_group.checkedId()

    def get_format(self) -> str:
        """'csv' or 'xlsx'"""
        return 'xlsx' if self.format_button_group.checkedId() == 1 else 'csv'


# ==================== Export Thread ====================

class ExportThread(QThread):
    """Thread for exporting data to CSV or XLSX."""
    progress = pyqtSignal(str)
    done = pyqtSignal(str)   # filepath
    error = pyqtSignal(str)

    def __init__(self, filepath: str, fmt: str, headers: List[str], rows: List[list]):
        super().__init__()
        self.filepath = filepath
        self.fmt = fmt
        self.headers = headers
        self.rows = rows

    def run(self):
        try:
            self.progress.emit("Экспорт данных...")
            if self.fmt == 'csv':
                with open(self.filepath, 'w', newline='', encoding='utf-8-sig') as f:
                    writer = csv.writer(f)
                    writer.writerow(self.headers)
                    writer.writerows(self.rows)
            elif self.fmt == 'xlsx':
                if openpyxl is None:
                    self.error.emit("openpyxl не установлен")
                    return
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "LemmaCheck Export"
                # Header row with bold font
                from openpyxl.styles import Font as XlFont
                bold = XlFont(bold=True)
                for col_idx, header in enumerate(self.headers, 1):
                    cell = ws.cell(row=1, column=col_idx, value=header)
                    cell.font = bold
                # Data rows
                for row_idx, row in enumerate(self.rows, 2):
                    for col_idx, value in enumerate(row, 1):
                        ws.cell(row=row_idx, column=col_idx, value=value)
                # Auto-width columns
                for col_idx, header in enumerate(self.headers, 1):
                    max_len = len(str(header))
                    for row in self.rows[:200]:
                        if col_idx - 1 < len(row):
                            max_len = max(max_len, len(str(row[col_idx - 1])))
                    col_letter = openpyxl.utils.get_column_letter(col_idx)
                    ws.column_dimensions[col_letter].width = min(max_len + 2, 60)
                wb.save(self.filepath)
            self.done.emit(self.filepath)
        except Exception as e:
            self.error.emit(str(e))


# ==================== Main Window ====================

class LemmaCheckApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.engine = LemmaSearchEngine()
        self.doc_paths: Dict[str, str] = {}
        self.result_doc_ids: List[str] = []
        self.indexing_thread: Optional[IndexingThread] = None
        self.search_thread: Optional[SearchThread] = None
        self.export_thread: Optional[ExportThread] = None
        # Stored data for export
        self.last_results: List[Tuple[str, int, str, List[int]]] = []
        self.last_query: str = ''
        self.last_kwic_data: List[Tuple[str, str, str, str, int]] = []
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("LemmaCheck - Полнотекстовый поиск по леммам")
        self.setGeometry(100, 100, 900, 700)
        self.setMinimumSize(800, 600)

        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setSpacing(10)
        layout.setContentsMargins(10, 10, 10, 10)

        # === Documents Section ===
        doc_group = QGroupBox("📁 Документы")
        doc_layout = QVBoxLayout(doc_group)

        btn_layout = QHBoxLayout()
        self.btn_add_files = QPushButton("Добавить файлы")
        self.btn_add_files.clicked.connect(self.add_files)
        btn_layout.addWidget(self.btn_add_files)

        self.btn_add_folder = QPushButton("Добавить папку")
        self.btn_add_folder.clicked.connect(self.add_folder)
        btn_layout.addWidget(self.btn_add_folder)

        self.btn_remove = QPushButton("Удалить")
        self.btn_remove.setStyleSheet("background-color: #e91e63; color: white;")
        self.btn_remove.clicked.connect(self.remove_selected)
        btn_layout.addWidget(self.btn_remove)

        self.btn_clear = QPushButton("Очистить")
        self.btn_clear.setStyleSheet("background-color: #e91e63; color: white;")
        self.btn_clear.clicked.connect(self.clear_all)
        btn_layout.addWidget(self.btn_clear)

        btn_layout.addStretch()

        self.btn_save = QPushButton("Сохранить индекс")
        self.btn_save.setStyleSheet("background-color: #ab47bc; color: white;")
        self.btn_save.clicked.connect(self.save_index)
        btn_layout.addWidget(self.btn_save)

        self.btn_load = QPushButton("Загрузить индекс")
        self.btn_load.setStyleSheet("background-color: #ab47bc; color: white;")
        self.btn_load.clicked.connect(self.load_index)
        btn_layout.addWidget(self.btn_load)

        doc_layout.addLayout(btn_layout)

        self.doc_list = QListWidget()
        self.doc_list.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        self.doc_list.setMaximumHeight(100)
        doc_layout.addWidget(self.doc_list)

        # Progress
        self.progress_widget = QWidget()
        progress_layout = QHBoxLayout(self.progress_widget)
        progress_layout.setContentsMargins(0, 0, 0, 0)
        self.progress_label = QLabel("")
        progress_layout.addWidget(self.progress_label)
        self.progress_bar = QProgressBar()
        progress_layout.addWidget(self.progress_bar)
        self.btn_cancel = QPushButton("Отмена")
        self.btn_cancel.clicked.connect(self.cancel_indexing)
        progress_layout.addWidget(self.btn_cancel)
        self.progress_widget.hide()
        doc_layout.addWidget(self.progress_widget)

        layout.addWidget(doc_group)

        # === Search Section ===
        search_group = QGroupBox("🔍 Поиск")
        search_layout = QHBoxLayout(search_group)
        
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Введите поисковый запрос...")
        self.search_input.setFont(QFont("Helvetica", 14))
        self.search_input.returnPressed.connect(self.search)
        search_layout.addWidget(self.search_input)

        self.btn_search = QPushButton("Искать")
        self.btn_search.setStyleSheet("background-color: #ec407a; color: white; font-size: 14px;")
        self.btn_search.clicked.connect(self.search)
        search_layout.addWidget(self.btn_search)

        self.kazakh_checkbox = QCheckBox("Казахский язык")
        self.kazakh_checkbox.setToolTip(
            "Все кириллические слова обрабатываются казахским стеммером\n"
            "(для слов без специфических казахских букв, напр. 'бала')"
        )
        self.kazakh_checkbox.toggled.connect(self.on_kazakh_toggled)
        search_layout.addWidget(self.kazakh_checkbox)

        layout.addWidget(search_group)

        # === KWIC & Export Section ===
        kwic_group = QGroupBox("📊 KWIC-конкорданс и экспорт")
        kwic_layout = QHBoxLayout(kwic_group)

        self.kwic_checkbox = QCheckBox("Включить KWIC")
        self.kwic_checkbox.setToolTip("Показать результаты в виде KWIC-таблицы (Key Word In Context)")
        kwic_layout.addWidget(self.kwic_checkbox)

        kwic_layout.addWidget(QLabel("Контекст:"))
        self.kwic_context_type = QComboBox()
        self.kwic_context_type.addItems(["±5 слов", "±10 слов", "Предложение"])
        self.kwic_context_type.setToolTip("Тип контекстного окна")
        kwic_layout.addWidget(self.kwic_context_type)

        kwic_layout.addWidget(QLabel("Фильтр:"))
        self.kwic_filter_input = QLineEdit()
        self.kwic_filter_input.setPlaceholderText("Слово в контексте...")
        self.kwic_filter_input.setMaximumWidth(160)
        self.kwic_filter_input.setToolTip(
            "Показать только строки, где в контексте (левом или правом) встречается это слово"
        )
        kwic_layout.addWidget(self.kwic_filter_input)

        kwic_layout.addStretch()

        self.btn_export = QPushButton("📥 Экспорт")
        self.btn_export.setStyleSheet("background-color: #ab47bc; color: white;")
        self.btn_export.setToolTip("Экспортировать результаты в CSV или XLSX")
        self.btn_export.clicked.connect(self.export_results)
        kwic_layout.addWidget(self.btn_export)

        layout.addWidget(kwic_group)

        # === Results Section ===
        results_group = QGroupBox("📋 Результаты")
        results_layout = QHBoxLayout(results_group)

        # Tabs for standard results and KWIC table
        self.results_tabs = QTabWidget()

        # Tab 1: Standard results
        self.results_text = QTextEdit()
        self.results_text.setReadOnly(True)
        self.results_text.setFont(QFont("Helvetica", 12))
        self.results_tabs.addTab(self.results_text, "Результаты")

        # Tab 2: KWIC table
        kwic_tab = QWidget()
        kwic_tab_layout = QVBoxLayout(kwic_tab)
        kwic_tab_layout.setContentsMargins(0, 0, 0, 0)
        self.kwic_table = QTableWidget()
        self.kwic_table.setColumnCount(4)
        self.kwic_table.setHorizontalHeaderLabels(["Левый контекст", "Ключевое слово", "Правый контекст", "Документ"])
        self.kwic_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self.kwic_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.kwic_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)
        self.kwic_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents)
        self.kwic_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        self.kwic_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.kwic_table.setAlternatingRowColors(True)
        self.kwic_table.setHorizontalScrollMode(QAbstractItemView.ScrollMode.ScrollPerPixel)
        self.kwic_table.setFont(QFont("Helvetica", 11))
        kwic_tab_layout.addWidget(self.kwic_table)

        kwic_btn_layout = QHBoxLayout()
        self.btn_copy_kwic = QPushButton("📋 Копировать таблицу")
        self.btn_copy_kwic.clicked.connect(self.copy_kwic_table)
        kwic_btn_layout.addWidget(self.btn_copy_kwic)
        kwic_btn_layout.addStretch()
        kwic_tab_layout.addLayout(kwic_btn_layout)

        self.results_tabs.addTab(kwic_tab, "KWIC-таблица")
        
        # Found words panel
        words_frame = QFrame()
        words_frame.setMaximumWidth(250)
        words_frame.setMinimumWidth(180)
        words_layout = QVBoxLayout(words_frame)
        words_layout.setContentsMargins(0, 0, 0, 0)
        
        words_label = QLabel("🔤 Найденные слова:")
        words_label.setFont(QFont("Helvetica", 11, QFont.Weight.Bold))
        words_layout.addWidget(words_label)
        
        self.found_words_text = QTextEdit()
        self.found_words_text.setReadOnly(True)
        self.found_words_text.setFont(QFont("Helvetica", 11))
        words_layout.addWidget(self.found_words_text)
        
        self.btn_copy_words = QPushButton("📋 Копировать список")
        self.btn_copy_words.clicked.connect(self.copy_found_words)
        words_layout.addWidget(self.btn_copy_words)
        
        # Use splitter for resizable panels
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.addWidget(self.results_tabs)
        splitter.addWidget(words_frame)
        splitter.setSizes([700, 200])
        
        results_layout.addWidget(splitter)

        layout.addWidget(results_group, stretch=1)

        # === Status Bar ===
        self.status_label = QLabel("Готово. Добавьте документы для начала работы.")
        layout.addWidget(self.status_label)

    def update_status(self):
        self.status_label.setText(
            f"Документов: {len(self.engine.documents)} | Индексировано слов: {self.engine.total_words}"
        )

    def add_files(self):
        files, _ = QFileDialog.getOpenFileNames(
            self, "Выберите документы", "",
            "Все поддерживаемые (*.docx *.doc *.pdf *.txt *.xlsx *.xls);;"
            "Word (*.docx *.doc);;PDF (*.pdf);;Text (*.txt);;Excel (*.xlsx *.xls)"
        )
        if files:
            self.index_files(files)

    def add_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку")
        if folder:
            files = []
            for root, _, filenames in os.walk(folder):
                for fn in filenames:
                    if fn.lower().endswith(('.docx', '.doc', '.pdf', '.txt', '.xlsx', '.xls')):
                        files.append(os.path.join(root, fn))
            if files:
                self.index_files(files)
            else:
                QMessageBox.information(self, "Информация", "Поддерживаемых документов не найдено.")

    def index_files(self, files: List[str]):
        if self.indexing_thread and self.indexing_thread.isRunning():
            QMessageBox.warning(self, "Внимание", "Индексация уже выполняется.")
            return

        self.progress_widget.show()
        self.progress_bar.setValue(0)
        self.set_buttons_enabled(False)

        self.indexing_thread = IndexingThread(files, self.engine)
        self.indexing_thread.progress.connect(self.on_progress)
        self.indexing_thread.finished_file.connect(self.on_file_indexed)
        self.indexing_thread.error.connect(self.on_indexing_error)
        self.indexing_thread.done.connect(self.on_indexing_done)
        self.indexing_thread.start()

    def on_progress(self, text: str, current: int, total: int):
        self.progress_label.setText(text)
        self.progress_bar.setValue(int(current / total * 100))

    def on_file_indexed(self, doc_id: str, filename: str):
        self.doc_paths[doc_id] = doc_id
        self.doc_list.addItem(filename)

    def on_indexing_error(self, error: str):
        QMessageBox.warning(self, "Ошибка", error)

    def on_indexing_done(self):
        self.progress_widget.hide()
        self.set_buttons_enabled(True)
        self.update_status()

    def on_kazakh_toggled(self, checked: bool):
        """Toggle Kazakh-only mode: re-index all documents with Kazakh stemmer."""
        self.engine.force_kazakh = checked
        self.engine._lemma_cache.clear()

        if not self.engine.documents:
            return

        # Re-index all documents with the new language setting
        stored = [(doc_id, info['text'], info['filename'])
                  for doc_id, info in self.engine.documents.items()]

        # Reset engine state
        self.engine.documents.clear()
        self.engine.inverted_index.clear()
        self.engine.doc_frequency.clear()
        self.engine.total_docs = 0
        self.engine.total_words = 0

        for doc_id, text, filename in stored:
            self.engine.add_document(doc_id, text, filename)

        self.update_status()
        mode = "казахский" if checked else "авто"
        self.status_label.setText(
            f"Режим: {mode} | Документов: {len(self.engine.documents)} | "
            f"Индексировано слов: {self.engine.total_words}"
        )

    def cancel_indexing(self):
        if self.indexing_thread:
            self.indexing_thread.cancelled = True

    def set_buttons_enabled(self, enabled: bool):
        for btn in [self.btn_add_files, self.btn_add_folder, self.btn_remove,
                    self.btn_clear, self.btn_save, self.btn_load, self.btn_search]:
            btn.setEnabled(enabled)

    def remove_selected(self):
        for item in self.doc_list.selectedItems():
            filename = item.text()
            for doc_id, doc in list(self.engine.documents.items()):
                if doc['filename'] == filename:
                    self.engine.remove_document(doc_id)
                    if doc_id in self.doc_paths:
                        del self.doc_paths[doc_id]
                    break
            self.doc_list.takeItem(self.doc_list.row(item))
        self.update_status()

    def clear_all(self):
        if not self.engine.documents:
            return
        if QMessageBox.question(self, "Подтверждение", "Удалить все документы?") == QMessageBox.StandardButton.Yes:
            self.engine = LemmaSearchEngine()
            self.doc_paths.clear()
            self.doc_list.clear()
            self.results_text.clear()
            self.found_words_text.clear()
            self.kwic_table.setRowCount(0)
            self.last_results = []
            self.last_query = ''
            self.last_kwic_data = []
            self.update_status()

    def save_index(self):
        if not self.engine.documents:
            QMessageBox.information(self, "Информация", "Нет документов для сохранения.")
            return
        filepath, _ = QFileDialog.getSaveFileName(self, "Сохранить индекс", "", "JSON (*.json)")
        if filepath:
            try:
                self.engine.save_index(filepath)
                paths_file = filepath.replace('.json', '_paths.json')
                with open(paths_file, 'w', encoding='utf-8') as f:
                    json.dump({'doc_paths': self.doc_paths}, f, ensure_ascii=False)
                QMessageBox.information(self, "Успешно", "Индекс сохранён.")
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", str(e))

    def load_index(self):
        filepath, _ = QFileDialog.getOpenFileName(self, "Загрузить индекс", "", "JSON (*.json)")
        if filepath:
            try:
                self.engine.load_index(filepath)
                paths_file = filepath.replace('.json', '_paths.json')
                if os.path.exists(paths_file):
                    with open(paths_file, 'r', encoding='utf-8') as f:
                        self.doc_paths = json.load(f).get('doc_paths', {})
                else:
                    self.doc_paths = {doc_id: doc_id for doc_id in self.engine.documents}
                self.doc_list.clear()
                for doc in self.engine.documents.values():
                    self.doc_list.addItem(doc['filename'])
                # Sync checkbox with loaded index setting
                self.kazakh_checkbox.blockSignals(True)
                self.kazakh_checkbox.setChecked(self.engine.force_kazakh)
                self.kazakh_checkbox.blockSignals(False)
                self.update_status()
                QMessageBox.information(self, "Успешно", "Индекс загружен.")
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", str(e))

    def search(self):
        query = self.search_input.text().strip()
        if not query:
            return
        if not self.engine.documents:
            QMessageBox.information(self, "Информация", "Сначала добавьте документы.")
            return
        
        # Show loading state
        self.results_text.clear()
        self.found_words_text.clear()
        self.results_text.setPlainText("⏳ Поиск...")
        self.btn_search.setEnabled(False)
        self.status_label.setText("Выполняется поиск...")
        QApplication.processEvents()
        
        # Determine KWIC settings
        kwic_enabled = self.kwic_checkbox.isChecked()
        kwic_ctx_idx = self.kwic_context_type.currentIndex()
        if kwic_ctx_idx == 0:
            kwic_context_type = 'words'
            kwic_context_words = 5
        elif kwic_ctx_idx == 1:
            kwic_context_type = 'words'
            kwic_context_words = 10
        else:
            kwic_context_type = 'sentence'
            kwic_context_words = 5
        kwic_filter = self.kwic_filter_input.text().strip()
        
        # Run search in thread
        self.search_thread = SearchThread(
            self.engine, query,
            kwic_enabled=kwic_enabled,
            kwic_context_type=kwic_context_type,
            kwic_context_words=kwic_context_words,
            kwic_filter=kwic_filter
        )
        self.search_thread.result_ready.connect(self.on_search_results)
        self.search_thread.kwic_ready.connect(self.on_kwic_results)
        self.search_thread.start()

    def on_search_results(self, results, query):
        """Handle search results from thread"""
        self.btn_search.setEnabled(True)
        self.last_results = results
        self.last_query = query
        self.results_text.clear()
        self.results_text.setPlainText("⏳ Обработка результатов...")
        QApplication.processEvents()
        
        self.display_results(results, query)
        self.update_status()

    def on_kwic_results(self, kwic_data):
        """Handle KWIC data from search thread"""
        self.last_kwic_data = kwic_data
        self._populate_kwic_table(kwic_data)
        self.results_tabs.setCurrentIndex(1)

    def display_results(self, results: List[Tuple[str, int, str, List[int]]], query: str):
        self.results_text.clear()
        self.found_words_text.clear()
        self.result_doc_ids = []

        if not results:
            self.results_text.setPlainText("Ничего не найдено.")
            return

        query_lemmas = self.engine.get_query_lemmas(query)
        query_lemmas_set = set(query_lemmas)
        is_phrase_search = len(query_lemmas) > 1
        
        # Collect found phrases/words per document with word counts
        phrases_by_doc: Dict[str, List[str]] = {}  # filename -> list of found phrases
        doc_word_counts: Dict[str, int] = {}  # filename -> total word count in doc
        
        cursor = self.results_text.textCursor()
        
        # Formats
        title_fmt = QTextCharFormat()
        title_fmt.setFontWeight(700)
        title_fmt.setForeground(QColor("#ad1457"))
        
        count_fmt = QTextCharFormat()
        count_fmt.setForeground(QColor("#8e24aa"))
        
        para_num_fmt = QTextCharFormat()
        para_num_fmt.setForeground(QColor("#ce93d8"))
        para_num_fmt.setFontWeight(600)
        
        # Soft highlight - light blue background, readable
        highlight_fmt = QTextCharFormat()
        highlight_fmt.setBackground(QColor("#fce4ec"))  # Light pink
        highlight_fmt.setForeground(QColor("#c2185b"))  # Deep pink text
        highlight_fmt.setFontWeight(600)
        
        normal_fmt = QTextCharFormat()
        
        total_results = len(results)

        for idx, (doc_id, count, filename, positions) in enumerate(results):
            # Update progress
            self.status_label.setText(f"Обработка документа {idx + 1} из {total_results}...")
            QApplication.processEvents()
            
            self.result_doc_ids.append(doc_id)
            phrases_by_doc[filename] = []
            doc_word_counts[filename] = self.engine.documents[doc_id]['word_count']
            
            full_text = self.engine.documents[doc_id]['text']
            
            if is_phrase_search:
                # For phrase search, extract the actual phrases found
                for pos in positions:
                    phrase = self._extract_phrase_at_position(full_text, pos, len(query_lemmas))
                    if phrase:
                        phrases_by_doc[filename].append(phrase)
                total_count = count  # count already contains phrase matches
            else:
                # For single word search, count all occurrences
                total_count = 0
                for match in re.finditer(WORD_PATTERN, full_text):
                    word = match.group()
                    lemma = self.engine._lemmatize_word(word)
                    if lemma in query_lemmas_set:
                        phrases_by_doc[filename].append(word.lower())
                        total_count += 1
            
            cursor.insertText(f"📄 {filename}", title_fmt)
            cursor.insertText(f"  [найдено: {total_count}]\n\n", count_fmt)
            
            # Get sentences with matches (1 before, match, 1 after)
            sentence_contexts = self.engine.get_sentences_with_matches(doc_id, positions)
            
            for match_num, context_text in sentence_contexts:
                cursor.insertText(f"[{match_num}] ", para_num_fmt)
                
                # Highlight matches in context
                if is_phrase_search:
                    self._highlight_phrases_in_context(cursor, context_text, query_lemmas, highlight_fmt, normal_fmt)
                else:
                    last_end = 0
                    for match in re.finditer(WORD_PATTERN, context_text):
                        word = match.group()
                        start = match.start()
                        
                        if start > last_end:
                            cursor.insertText(context_text[last_end:start], normal_fmt)
                        
                        lemma = self.engine._lemmatize_word(word)
                        if lemma in query_lemmas_set:
                            cursor.insertText(word, highlight_fmt)
                        else:
                            cursor.insertText(word, normal_fmt)
                        
                        last_end = match.end()
                    
                    if last_end < len(context_text):
                        cursor.insertText(context_text[last_end:], normal_fmt)
                
                cursor.insertText("\n\n", normal_fmt)
            
            cursor.insertText("─" * 70 + "\n\n", normal_fmt)

        cursor.insertText(f"Найдено документов: {len(results)}", normal_fmt)
        
        # Populate found phrases panel grouped by document
        words_cursor = self.found_words_text.textCursor()
        
        doc_title_fmt = QTextCharFormat()
        doc_title_fmt.setFontWeight(700)
        doc_title_fmt.setForeground(QColor("#ad1457"))
        
        total_fmt = QTextCharFormat()
        total_fmt.setFontWeight(700)
        total_fmt.setForeground(QColor("#c2185b"))
        
        stats_fmt = QTextCharFormat()
        stats_fmt.setForeground(QColor("#6a1b4d"))
        
        variation_fmt = QTextCharFormat()
        variation_fmt.setForeground(QColor("#8e24aa"))  # Purple for variations count
        variation_fmt.setFontWeight(600)
        
        percent_fmt = QTextCharFormat()
        percent_fmt.setForeground(QColor("#d81b60"))  # Hot pink for percentage
        
        ip10k_fmt = QTextCharFormat()
        ip10k_fmt.setForeground(QColor("#ab47bc"))  # Purple for IP10K
        
        ipm_fmt = QTextCharFormat()
        ipm_fmt.setForeground(QColor("#ec407a"))  # Pink for IPM
        
        word_fmt = QTextCharFormat()
        
        grand_total = 0
        grand_variations = 0
        total_words_all_docs = 0
        
        for filename, phrases in phrases_by_doc.items():
            if phrases:
                # Count unique phrases
                phrase_counts = defaultdict(int)
                for phrase in phrases:
                    phrase_counts[phrase.lower()] += 1
                
                doc_total = sum(phrase_counts.values())
                doc_variations = len(phrase_counts)
                doc_word_count = doc_word_counts.get(filename, 1)
                doc_percent = (doc_total / doc_word_count * 100) if doc_word_count > 0 else 0
                doc_iptт = (doc_total / doc_word_count * 10_000) if doc_word_count > 0 else 0  # per 10K
                doc_ipm = (doc_total / doc_word_count * 1_000_000) if doc_word_count > 0 else 0  # per 1M
                
                grand_total += doc_total
                grand_variations += doc_variations
                total_words_all_docs += doc_word_count
                
                words_cursor.insertText(f"📄 {filename}\n", doc_title_fmt)
                words_cursor.insertText(f"   Всего: {doc_total} | ", stats_fmt)
                words_cursor.insertText(f"Вариаций: {doc_variations}\n", variation_fmt)
                words_cursor.insertText(f"   Частота: {doc_percent:.2f}% | ", percent_fmt)
                words_cursor.insertText(f"IP10K: {doc_iptт:.2f}", ip10k_fmt)
                words_cursor.insertText(f" | ", stats_fmt)
                words_cursor.insertText(f"IPM: {doc_ipm:.1f}\n", ipm_fmt)
                
                sorted_phrases = sorted(phrase_counts.items(), key=lambda x: (-x[1], x[0]))
                for phrase, cnt in sorted_phrases:
                    phrase_percent = (cnt / doc_word_count * 100) if doc_word_count > 0 else 0
                    phrase_iptт = (cnt / doc_word_count * 10_000) if doc_word_count > 0 else 0
                    phrase_ipm = (cnt / doc_word_count * 1_000_000) if doc_word_count > 0 else 0
                    words_cursor.insertText(f"  • {phrase} ({cnt}) ", word_fmt)
                    words_cursor.insertText(f"— {phrase_percent:.2f}% ", percent_fmt)
                    words_cursor.insertText(f"[{phrase_iptт:.1f}", ip10k_fmt)
                    words_cursor.insertText(f" | ", stats_fmt)
                    words_cursor.insertText(f"{phrase_ipm:.0f}]\n", ipm_fmt)
                words_cursor.insertText("\n", word_fmt)
        
        # Calculate grand total percentage, IP10K and IPM
        grand_percent = (grand_total / total_words_all_docs * 100) if total_words_all_docs > 0 else 0
        grand_iptт = (grand_total / total_words_all_docs * 10_000) if total_words_all_docs > 0 else 0
        grand_ipm = (grand_total / total_words_all_docs * 1_000_000) if total_words_all_docs > 0 else 0
        
        # Add grand total
        words_cursor.insertText("═" * 25 + "\n", word_fmt)
        words_cursor.insertText(f"📊 ИТОГО\n", total_fmt)
        words_cursor.insertText(f"   Совпадений: {grand_total} ", stats_fmt)
        words_cursor.insertText(f"({grand_percent:.2f}%)\n", percent_fmt)
        words_cursor.insertText(f"   IP10K: ", stats_fmt)
        words_cursor.insertText(f"{grand_iptт:.2f}\n", ip10k_fmt)
        words_cursor.insertText(f"   IPM: ", stats_fmt)
        words_cursor.insertText(f"{grand_ipm:.1f}\n", ipm_fmt)
        words_cursor.insertText(f"   Вариаций: {grand_variations}\n", variation_fmt)
        words_cursor.insertText(f"   Документов: {len(results)}\n", stats_fmt)
        words_cursor.insertText(f"   Слов всего: {total_words_all_docs:,}".replace(',', ' '), stats_fmt)
        
        self.status_label.setText(f"Найдено {grand_total} совпадений ({grand_percent:.2f}%, IP10K: {grand_iptт:.1f}) в {len(results)} документах")

        # Clear KWIC table if KWIC is not enabled (KWIC is built via on_kwic_results signal)
        if not self.kwic_checkbox.isChecked():
            self.kwic_table.setRowCount(0)
            self.last_kwic_data = []

    def _populate_kwic_table(self, kwic_data: List[Tuple[str, str, str, str, int]]):
        """Populate KWIC table from pre-computed data (built in SearchThread)."""
        self.kwic_table.setRowCount(len(kwic_data))
        kw_font = QFont("Helvetica", 11, QFont.Weight.Bold)
        right_align = Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter
        left_align = Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter
        center_align = Qt.AlignmentFlag.AlignCenter

        for row_idx, (fname, left, kw, right, pos) in enumerate(kwic_data):
            # Left context (right-aligned)
            item_left = QTableWidgetItem(left)
            item_left.setTextAlignment(right_align)
            self.kwic_table.setItem(row_idx, 0, item_left)
            # Keyword (bold, centered)
            item_kw = QTableWidgetItem(kw)
            item_kw.setFont(kw_font)
            item_kw.setTextAlignment(center_align)
            item_kw.setForeground(QColor("#880e4f"))
            item_kw.setBackground(QColor("#fce4ec"))
            self.kwic_table.setItem(row_idx, 1, item_kw)
            # Right context (left-aligned)
            item_right = QTableWidgetItem(right)
            item_right.setTextAlignment(left_align)
            self.kwic_table.setItem(row_idx, 2, item_right)
            # Document
            item_doc = QTableWidgetItem(fname)
            item_doc.setTextAlignment(left_align)
            self.kwic_table.setItem(row_idx, 3, item_doc)

        self.kwic_table.resizeRowsToContents()
        self.status_label.setText(
            f"KWIC: {len(kwic_data)} конкорданс{'ов' if len(kwic_data) != 1 else ''}"
        )

    def copy_kwic_table(self):
        """Copy KWIC table to clipboard as tab-separated text."""
        rows = self.kwic_table.rowCount()
        if rows == 0:
            return
        lines = ["Левый контекст\tКлючевое слово\tПравый контекст\tДокумент"]
        for r in range(rows):
            cols = []
            for c in range(4):
                item = self.kwic_table.item(r, c)
                cols.append(item.text() if item else '')
            lines.append('\t'.join(cols))
        text = '\n'.join(lines)
        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        self.status_label.setText(f"KWIC-таблица скопирована ({rows} строк)")

    def _extract_phrase_at_position(self, text: str, start_pos: int, word_count: int) -> str:
        """Extract a phrase starting at given position with given number of words"""
        words = []
        for match in re.finditer(WORD_PATTERN, text[start_pos:]):
            words.append(match.group())
            if len(words) >= word_count:
                break
        return ' '.join(words) if len(words) == word_count else ''

    def _highlight_phrases_in_context(self, cursor, context: str, query_lemmas: List[str], 
                                       highlight_fmt, normal_fmt):
        """Highlight phrase matches in context text"""
        # Find all word positions in context
        word_matches = list(re.finditer(WORD_PATTERN, context))
        
        if not word_matches:
            cursor.insertText(context, normal_fmt)
            return
        
        # Find phrase matches (sequences of words matching query lemmas)
        highlight_ranges = []
        i = 0
        while i <= len(word_matches) - len(query_lemmas):
            matched = True
            for j, query_lemma in enumerate(query_lemmas):
                word = word_matches[i + j].group()
                word_lemma = self.engine._lemmatize_word(word)
                if word_lemma != query_lemma:
                    matched = False
                    break
            
            if matched:
                start = word_matches[i].start()
                end = word_matches[i + len(query_lemmas) - 1].end()
                highlight_ranges.append((start, end))
                i += len(query_lemmas)  # Skip past matched phrase
            else:
                i += 1
        
        # Output text with highlights
        last_end = 0
        for start, end in highlight_ranges:
            if start > last_end:
                cursor.insertText(context[last_end:start], normal_fmt)
            cursor.insertText(context[start:end], highlight_fmt)
            last_end = end
        
        if last_end < len(context):
            cursor.insertText(context[last_end:], normal_fmt)

    # ==================== Export Methods ====================

    def export_results(self):
        """Show export dialog and perform export."""
        if not self.last_results:
            QMessageBox.information(self, "Информация", "Сначала выполните поиск.")
            return

        has_kwic = len(self.last_kwic_data) > 0
        dialog = ExportDialog(self, has_kwic=has_kwic)

        if dialog.exec() != QDialog.DialogCode.Accepted:
            return

        export_type = dialog.get_export_type()
        fmt = dialog.get_format()

        # Build headers and rows
        if export_type == 0:
            headers, rows = self._compute_results_table()
        elif export_type == 1:
            headers, rows = self._compute_concordance_table()
        else:
            headers, rows = self._compute_summary()

        if not rows:
            QMessageBox.information(self, "Информация", "Нет данных для экспорта.")
            return

        # Generate default filename
        now = datetime.now()
        type_label = ['results', 'concordance', 'summary'][export_type]
        default_name = f"lemmaCheck_export_{type_label}_{now.strftime('%Y%m%d_%H%M')}.{fmt}"

        ext_filter = "CSV (*.csv)" if fmt == 'csv' else "Excel (*.xlsx)"
        filepath, _ = QFileDialog.getSaveFileName(
            self, "Экспорт данных", default_name, ext_filter
        )
        if not filepath:
            return

        # Run export in background thread
        self.export_thread = ExportThread(filepath, fmt, headers, rows)
        self.export_thread.progress.connect(lambda msg: self.status_label.setText(msg))
        self.export_thread.done.connect(self._on_export_done)
        self.export_thread.error.connect(self._on_export_error)
        self.export_thread.start()

    def _on_export_done(self, filepath: str):
        """Handle export completion."""
        filename = os.path.basename(filepath)
        self.status_label.setText(f"Экспорт завершён: {filename}")
        QMessageBox.information(self, "Экспорт", f"Данные успешно экспортированы:\n{filepath}")

    def _on_export_error(self, error: str):
        """Handle export error."""
        self.status_label.setText("Ошибка экспорта")
        QMessageBox.critical(self, "Ошибка экспорта", error)

    def _compute_results_table(self) -> Tuple[List[str], List[list]]:
        """Compute results table data for export.
        Columns: Лемма, Словоформа, Документ, Абс.частота, %, IP10K, IPM, TF-IDF
        """
        headers = ["Лемма", "Словоформа", "Документ", "Абс. частота", "%", "IP10K", "IPM", "TF-IDF"]
        rows: List[list] = []

        if not self.last_results or not self.last_query:
            return headers, rows

        query_lemmas = self.engine.get_query_lemmas(self.last_query)
        query_lemmas_set = set(query_lemmas)
        is_phrase = len(query_lemmas) > 1
        lemma_str = ' '.join(query_lemmas)

        for doc_id, count, filename, positions in self.last_results:
            if doc_id not in self.engine.documents:
                continue
            full_text = self.engine.documents[doc_id]['text']
            word_count = self.engine.documents[doc_id]['word_count']

            # Collect word forms and their frequencies
            word_forms: Dict[str, int] = defaultdict(int)
            if is_phrase:
                for pos in positions:
                    phrase = self._extract_phrase_at_position(full_text, pos, len(query_lemmas))
                    if phrase:
                        word_forms[phrase.lower()] += 1
            else:
                for match in re.finditer(WORD_PATTERN, full_text):
                    word = match.group()
                    lemma = self.engine._lemmatize_word(word)
                    if lemma in query_lemmas_set:
                        word_forms[word.lower()] += 1

            # TF-IDF sum for query lemmas in this doc
            tfidf_sum = sum(self.engine._calculate_tfidf(l, doc_id) for l in query_lemmas_set)

            for form, freq in sorted(word_forms.items(), key=lambda x: -x[1]):
                pct = (freq / word_count * 100) if word_count > 0 else 0
                ip10k = (freq / word_count * 10_000) if word_count > 0 else 0
                ipm = (freq / word_count * 1_000_000) if word_count > 0 else 0
                rows.append([
                    lemma_str, form, filename, freq,
                    round(pct, 4), round(ip10k, 2), round(ipm, 1), round(tfidf_sum, 6)
                ])

        return headers, rows

    def _compute_concordance_table(self) -> Tuple[List[str], List[list]]:
        """Compute concordance (KWIC) table data for export."""
        headers = ["Левый контекст", "Ключевое слово", "Правый контекст", "Документ", "Позиция"]
        rows: List[list] = []

        for fname, left, kw, right, pos in self.last_kwic_data:
            rows.append([left, kw, right, fname, pos])

        return headers, rows

    def _compute_summary(self) -> Tuple[List[str], List[list]]:
        """Compute summary report data for export."""
        headers = ["Параметр", "Значение"]
        now = datetime.now()
        query = self.last_query or ''
        lang_mode = "Казахский" if self.engine.force_kazakh else "Авто (Рус/Англ/Каз)"
        total_docs = len(self.engine.documents)
        total_words = self.engine.total_words
        match_docs = len(self.last_results) if self.last_results else 0
        total_matches = sum(r[1] for r in self.last_results) if self.last_results else 0

        # Corpus-wide frequency stats
        grand_pct = (total_matches / total_words * 100) if total_words > 0 else 0
        grand_ip10k = (total_matches / total_words * 10_000) if total_words > 0 else 0
        grand_ipm = (total_matches / total_words * 1_000_000) if total_words > 0 else 0

        rows: List[list] = [
            ["Дата и время", now.strftime("%Y-%m-%d %H:%M:%S")],
            ["Поисковый запрос", query],
            ["Языковой режим", lang_mode],
            ["Документов в корпусе", total_docs],
            ["Слов в корпусе", total_words],
            ["Документов с совпадениями", match_docs],
            ["Всего совпадений", total_matches],
            ["Частота (%)", round(grand_pct, 4)],
            ["IP10K", round(grand_ip10k, 2)],
            ["IPM", round(grand_ipm, 1)],
            ["KWIC-строк", len(self.last_kwic_data)],
            ["Версия LemmaCheck", "1.0"],
        ]
        return headers, rows

    def open_file(self, filepath: str):
        try:
            if platform.system() == 'Darwin':
                subprocess.run(['open', filepath], check=True)
            elif platform.system() == 'Windows':
                os.startfile(filepath)
            else:
                subprocess.run(['xdg-open', filepath], check=True)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось открыть файл: {e}")

    def copy_found_words(self):
        """Copy all found words to clipboard"""
        text = self.found_words_text.toPlainText()
        
        if text.strip():
            clipboard = QApplication.clipboard()
            clipboard.setText(text)
            lines = [l for l in text.split('\n') if l.strip()]
            self.status_label.setText(f"Скопировано {len(lines)} строк в буфер обмена")


def main():
    app = QApplication(sys.argv)
    app.setStyle('Fusion')

    # Barbie pink palette
    palette = QPalette()
    palette.setColor(QPalette.ColorRole.Window, QColor("#fff0f5"))          # Lavender blush background
    palette.setColor(QPalette.ColorRole.WindowText, QColor("#4a1942"))      # Deep plum text
    palette.setColor(QPalette.ColorRole.Base, QColor("#ffffff"))             # White input fields
    palette.setColor(QPalette.ColorRole.AlternateBase, QColor("#fce4ec"))   # Light pink alternating rows
    palette.setColor(QPalette.ColorRole.ToolTipBase, QColor("#f8bbd0"))     # Pink tooltips
    palette.setColor(QPalette.ColorRole.ToolTipText, QColor("#4a1942"))     # Deep plum tooltip text
    palette.setColor(QPalette.ColorRole.Text, QColor("#4a1942"))            # Deep plum text
    palette.setColor(QPalette.ColorRole.Button, QColor("#f8bbd0"))          # Pink buttons
    palette.setColor(QPalette.ColorRole.ButtonText, QColor("#880e4f"))      # Dark pink button text
    palette.setColor(QPalette.ColorRole.BrightText, QColor("#ff4081"))      # Hot pink bright text
    palette.setColor(QPalette.ColorRole.Link, QColor("#d81b60"))            # Pink links
    palette.setColor(QPalette.ColorRole.Highlight, QColor("#f06292"))       # Pink selection
    palette.setColor(QPalette.ColorRole.HighlightedText, QColor("#ffffff")) # White on selection
    palette.setColor(QPalette.ColorRole.PlaceholderText, QColor("#ce93d8")) # Soft purple placeholder
    app.setPalette(palette)

    # Global stylesheet for Barbie pink theme
    app.setStyleSheet("""
        QGroupBox {
            border: 2px solid #f48fb1;
            border-radius: 8px;
            margin-top: 10px;
            padding-top: 14px;
            font-weight: bold;
            color: #880e4f;
        }
        QGroupBox::title {
            subcontrol-origin: margin;
            left: 12px;
            padding: 0 6px;
            color: #ad1457;
        }
        QPushButton {
            background-color: #f48fb1;
            color: white;
            border: none;
            border-radius: 6px;
            padding: 6px 14px;
            font-weight: bold;
        }
        QPushButton:hover {
            background-color: #f06292;
        }
        QPushButton:pressed {
            background-color: #ec407a;
        }
        QPushButton:disabled {
            background-color: #f8bbd0;
            color: #e0e0e0;
        }
        QLineEdit, QComboBox {
            border: 2px solid #f48fb1;
            border-radius: 6px;
            padding: 4px 8px;
            background: white;
        }
        QLineEdit:focus, QComboBox:focus {
            border-color: #ec407a;
        }
        QListWidget {
            border: 2px solid #f8bbd0;
            border-radius: 6px;
            background: white;
        }
        QTextEdit {
            border: 2px solid #f8bbd0;
            border-radius: 6px;
            background: white;
        }
        QTableWidget {
            border: 2px solid #f8bbd0;
            border-radius: 6px;
            gridline-color: #fce4ec;
            background: white;
        }
        QHeaderView::section {
            background-color: #f8bbd0;
            color: #880e4f;
            padding: 4px;
            border: 1px solid #f48fb1;
            font-weight: bold;
        }
        QTabWidget::pane {
            border: 2px solid #f48fb1;
            border-radius: 6px;
            background: white;
        }
        QTabBar::tab {
            background: #fce4ec;
            color: #880e4f;
            padding: 6px 16px;
            border-top-left-radius: 6px;
            border-top-right-radius: 6px;
            margin-right: 2px;
        }
        QTabBar::tab:selected {
            background: #f48fb1;
            color: white;
        }
        QTabBar::tab:hover {
            background: #f06292;
            color: white;
        }
        QProgressBar {
            border: 2px solid #f48fb1;
            border-radius: 6px;
            text-align: center;
            background: #fce4ec;
            color: #880e4f;
        }
        QProgressBar::chunk {
            background-color: #f06292;
            border-radius: 4px;
        }
        QCheckBox {
            color: #880e4f;
            spacing: 6px;
        }
        QCheckBox::indicator:checked {
            background-color: #ec407a;
            border: 2px solid #ad1457;
            border-radius: 3px;
        }
        QCheckBox::indicator:unchecked {
            background-color: white;
            border: 2px solid #f48fb1;
            border-radius: 3px;
        }
        QLabel {
            color: #4a1942;
        }
        QScrollBar:vertical {
            background: #fce4ec;
            width: 12px;
            border-radius: 6px;
        }
        QScrollBar::handle:vertical {
            background: #f48fb1;
            border-radius: 6px;
            min-height: 20px;
        }
        QScrollBar::handle:vertical:hover {
            background: #f06292;
        }
    """)

    window = LemmaCheckApp()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
